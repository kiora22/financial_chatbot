"""
Microsoft Word document processor.
Handles .docx files using python-docx.
"""

import os
from typing import Dict, Any, Optional
from docx import Document

from document_processor.processors.base_processor import BaseDocumentProcessor
from config.logging_config import setup_logging

# Set up logging
logger = setup_logging("docx_processor")


class DocxProcessor(BaseDocumentProcessor):
    """Processor for Microsoft Word documents (.docx)."""
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Extracted text
        """
        try:
            logger.info(f"Extracting text from Word document: {file_path}")
            
            # Open the Word document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            # Join all text
            text = "\n\n".join(paragraphs)
            
            logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {str(e)}")
            return ""
    
    async def extract_metadata(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Extract metadata from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Extracted metadata
        """
        try:
            # Get file stats
            stat_info = os.stat(file_path)
            
            # Initialize metadata with file stats
            metadata = {
                "file_size_bytes": stat_info.st_size,
                "created_at": stat_info.st_ctime,
                "modified_at": stat_info.st_mtime,
                "file_extension": "docx"
            }
            
            # Extract Word-specific metadata
            doc = Document(file_path)
            
            # Get core properties
            core_props = doc.core_properties
            
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["doc_created_at"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["doc_modified_at"] = core_props.modified.isoformat()
            if core_props.comments:
                metadata["comments"] = core_props.comments
            
            # Count paragraphs, sections and tables
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["section_count"] = len(doc.sections)
            metadata["table_count"] = len(doc.tables)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from Word document {file_path}: {str(e)}")
            return None